#!/usr/bin/python
# -*- coding: utf-8 -*-
#__author__:"watalo"
# @Time: 2020/4/13 21:25
# @Site    : 
# @File    : main.py
# @Software: PyCharm

import os
from tinydb import TinyDB, Query
from docx.shared import Inches
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from far import *

'''
    直接使用get_docx()生成docx文件
'''

def get_docx(name, output_path):
    conf = Conf(name)
    doc = Document()

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 文档标题
    title = doc.add_heading('',level=0).add_run('{}财务分析报告'.format(conf.name))
    title.font.name = 'Times New Roman'
    title._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    # 1.数据
    bold(doc, conf.header.h1,2)
    financial_sheet(conf, doc)
    doc.add_paragraph(':::::::请调整成自己喜欢的表格样式::::::')
    # 2.分析
    if conf.data_type == 'normal':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.normal.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.normal.s2d1.format(**conf.para.s2d1))
        doc.add_paragraph(conf.normal.s2d2.format(**conf.para.s2d2))
        p3 = bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.normal.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.normal.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.normal.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.normal.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h2s7,3)
        big_change_sheet(conf, doc)
        bold(doc, conf.header.h3,3)
        items_detail(conf, doc, 'year1')
    elif conf.data_type == 'no_3year':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.no_year3.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.no_year3.s2d1.format(**conf.para.s2d1))
        p3 = doc.add_paragraph(conf.no_year3.s2d2.format(**conf.para.s2d2))
        bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.no_year3.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.no_year3.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.no_year3.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.no_year3.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h2s7, 3)
        big_change_sheet(conf, doc)
        bold(doc, conf.header.h3, 2)
        items_detail(conf, doc, 'year1')
    elif conf.data_type == 'no_year2':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.no_year2.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.no_year2.s2d1.format(**conf.para.s2d1))
        p3 = doc.add_paragraph(conf.no_year2.s2d2.format(**conf.para.s2d2))
        bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.no_year2.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.no_year2.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.no_year2.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.no_year2.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h2s7, 3)
        big_change_sheet(conf, doc)
        bold(doc, conf.header.h3, 2)
        items_detail(conf, doc, 'year1')
    elif conf.data_type == 'no_year1':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.no_year1.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.no_year1.s2d1.format(**conf.para.s2d1))
        p3 = doc.add_paragraph(conf.no_year1.s2d2.format(**conf.para.s2d2))
        bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.no_year1.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.no_year1.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.no_year1.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.no_year1.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h3, 2)
        items_detail(conf, doc, 'month')
    elif conf.data_type == 'all_years':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.all_years.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.all_years.s2d1.format(**conf.para.s2d1))
        p3 = doc.add_paragraph(conf.all_years.s2d2.format(**conf.para.s2d2))
        bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.all_years.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.all_years.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.all_years.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.all_years.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h3,2)
        items_detail(conf, doc, 'year1')
    elif conf.data_type == 'two_years':
        bold(doc, conf.header.h2, 2)
        bold(doc, conf.header.h2s1, 3)
        p1 = doc.add_paragraph(conf.two_years.s1d1)
        bold(doc, conf.header.h2s2, 3)
        p2 = doc.add_paragraph(conf.two_years.s2d1.format(**conf.para.s2d1))
        p3 = doc.add_paragraph(conf.two_years.s2d2.format(**conf.para.s2d2))
        bold(doc, conf.header.h2s3, 3)
        p4 = doc.add_paragraph(conf.two_years.s2d3.format(**conf.para.s2d3))
        bold(doc, conf.header.h2s4, 3)
        p5 = doc.add_paragraph(conf.two_years.s2d4.format(**conf.para.s2d4))
        bold(doc, conf.header.h2s5, 3)
        p6 = doc.add_paragraph(conf.two_years.s2d5.format(**conf.para.s2d5))
        bold(doc, conf.header.h2s6, 3)
        p7 = doc.add_paragraph(conf.two_years.s2d6.format(**conf.para.s2d6))
        bold(doc, conf.header.h3,2)
        items_detail(conf, doc, 'year1')
    else:
        pass
    
    for paragraph in [p1,p2,p3,p4,p5,p6,p7]:
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)

    # 3.讨米文案
    doc.add_picture('img/taomi.png', width=Inches(2.25))
    p = doc.add_paragraph()
    run = p.add_run('后期开发计划是增加各个行业的行业分析，扫码赞赏可以加快开发速度哦，感谢您的支持！')
    run.bold = True
    # 4.保存
    doc.save(output_path)

#------------------------------配置类：调用其他类-------------------------
class Conf(object):

    def __init__(self, name):
        self.name = name
        self.db_path = _config.Path.db
        self.db_file_path = '/'.join([self.db_path, os.listdir(self.db_path)[0]])
        self.db = TinyDB(self.db_file_path)
        self.table = self.db.table('without_nodata')
        self.table_for_print = self.db.table('for_print')
        # para类的实例初始化
        self.data_type = _para.dict_().data_type
        self.para = _para.dict_()
        self.init_para()
        # text类的实例
        self.normal = _text.normal()
        self.no_year3 = _text.no_year3()
        self.no_year2 = _text.no_year2()
        self.no_year1 = _text.no_year1()
        self.all_years = _text.all_years()
        self.two_years = _text.two_years()
        self.header = _text.header()

    def init_para(self):
        self.para.set()
        self.para.analysis_s2d1()
        self.para.analysis_s2d2()
        self.para.analysis_s2d3()
        self.para.analysis_s2d5()

# -----------------------------需要用到的函数-----------------------------

def bold(doc, text, level):
    run = doc.add_heading('', level=level).add_run(text)
    run.font.name = 'Times Nwe Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def data(conf_obj,item, key):
    Q = Query()
    return conf_obj.table.get(Q.items == item)[key]

# -----------------------------------------------------------------------
def financial_sheet(conf_obj,doc):
    Q = Query()
    item_list = []  # 表头列表
    for item in conf_obj.table_for_print.all():
        item_list.append(item['items'])
    data_list = []  # 数据列表
    for data in conf_obj.table_for_print.search(Q.items == '项目'):
        for i, it in enumerate(data):
            if type(it) == str:
                data_list.append(it)
    data_list = data_list[:5]
# ------------------------------------------------------------------------
    xratio = [
        '资产负债率', '费用收入比', '流动资产占比',
        '毛利率', '净利润率', '刚兑占比', '短债占比',
        '总资产收益率(ROA)', '净资产收益率(ROE)',
    ]  # 百分数形式的财务指标 eg:‘11.11%’
    table = doc.add_table(rows=len(conf_obj.table_for_print.all()),
                          cols=5,
                          style="Light Grid")
# ------------------------------------------------------------------------
    for irows, item in enumerate(item_list):
        for icols, data in enumerate(data_list):
            num = conf_obj.table_for_print.get(Q.items == item)[data]
            if isinstance(num, float):
                if item in xratio:
                    table.cell(irows, icols).text = '{:.2%}'.format(num)
                else:
                    table.cell(irows, icols).text = '{:,.2f}'.format(num)
            else:
                table.cell(irows, icols).text = str(num)
# ------------------------------------------------------------------------

def big_change_items(conf_obj):
    '''
    不直接使用，只在big_change_sheet()中被调用。
    conf.data_type == 'no_year1'时不能是使用本函数
    找出最近一期比上一期变化超过30%的科目
    :return:
    '''
    Q = Query()
    Item = conf_obj.table.search((Q.type == '流动资产') | (Q.type == '非流动资产') | (Q.type == '流动负债') | (Q.type == '非流动负债'))
    list1 = []
    list2 = []
    big_change_items = []
    if conf_obj.data_type == 'all_years':
        for dict in Item:
            if dict['year1'] == 0 and dict['month'] != 0:
                list1.append(dict)
            elif dict['year1'] == 0 and dict['month'] == 0:
                pass
            else:
                if (dict['month'] - dict['year1']) / dict['year1'] < 0.3 \
                        and (dict['month'] - dict['year1']) / dict['year1'] > -0.3:
                    pass
                else:
                    list2.append(dict)
        list_sorted_dict = sorted(list2, key=lambda x: x['month'] - x['year1'], reverse=True)
        # 组成当期纯变化和超过30%变化的列表
        for dict in list_sorted_dict:
            list1.append(dict)
        # 形成科目名称列表
        for dict in list1:
            big_change_items.append(dict['items'])
        return big_change_items
    else:
        for dict in Item:
            if dict['year1'] == 0 and dict['year2'] != 0:
                list1.append(dict)
            elif dict['year1'] == 0 and dict['year2'] == 0:
                pass
            else:
                if (dict['year1']-dict['year2'])/dict['year2'] < 0.3 \
                        and (dict['year1']-dict['year2'])/dict['year2'] > -0.3:
                    pass
                else:
                    list2.append(dict)
        list_sorted_dict = sorted(list2, key=lambda x: x['year1']-x['year2'], reverse=True)
        #组成当期纯变化和超过30%变化的列表
        for dict in list_sorted_dict:
            list1.append(dict)
        # 形成科目名称列表
        for dict in list1:
            big_change_items.append(dict['items'])
        return big_change_items

def big_change_sheet(conf_obj,doc):
    table_change_item = ['科目名称', '当期值', '较年初变化', '变化率', '变化情况']
    table_change = doc.add_table(rows=len(big_change_items(conf_obj)) + 1,
                                 cols=5,
                                 style='Light Grid')
    for i in range(5):
        cell = table_change.cell(0, i)
        cell.text = table_change_item[i]
    if conf_obj.data_type == 'all_years':
        for i, item in enumerate(big_change_items(conf_obj)):
            table_change.cell(i + 1, 0).text = data(conf_obj,item,'items')  # 科目名称
            table_change.cell(i + 1, 1).text = '{:,.2f}'.format(data(conf_obj,item,'year1'))  # 当期值
            table_change.cell(i + 1, 2).text = '{:,.2f}'.format(
                data(conf_obj,item,'year1') - data(conf_obj,item,'year2'))  # 变化值
            if data(conf_obj,item,'year2') != 0:
                ratio = data(conf_obj,item,'year1') - data(conf_obj, item, 'year2') / data(conf_obj, item, 'year2')
                table_change.cell(i + 1, 3).text = '{:.2%}'.format(ratio)  # 变化率
            else:
                if data(conf_obj, item, 'year1') > 0:
                    table_change.cell(i + 1, 3).text = '当期净增加'
                else:
                    table_change.cell(i + 1, 3).text = '当期净减少'
    else:
        for i, item in enumerate(big_change_items(conf_obj)):
            table_change.cell(i + 1, 0).text = data(conf_obj,item,'items')  # 科目名称
            table_change.cell(i + 1, 1).text = '{:,.2f}'.format(data(conf_obj,item,'month'))  # 当期值
            table_change.cell(i + 1, 2).text = '{:,.2f}'.format(data(conf_obj,item,'month') - data(conf_obj,item, 'year1'))  # 变化值
            if data(conf_obj,item,'year1') != 0:
                ratio = (data(conf_obj,item,'month') - data(conf_obj,item,'year1')) / data(conf_obj,item,'year1')
                table_change.cell(i + 1, 3).text = '{:.2%}'.format(ratio)  # 变化率
            else:
                if data(conf_obj,item,'month') > 0:
                    table_change.cell(i + 1, 3).text = '当期净增加'
                else:
                    table_change.cell(i + 1, 3).text = '当期净减少'

def item_table(doc, col1, col2, col3, col4):
    """
    生成固定格式的表格
        1、科目明细，2、账龄明细，3、等等
    @param col1: 序号
    @param col2: 名称
    @param col3: 金额
    @param col4: 等等
    @return: 空白表格
    """
    table_regular = [col1, col2, col3, col4]
    b_c_table = doc.add_table(rows=7,
                              cols=4,
                              style='Light Grid')
    for i in range(4):
        cell = b_c_table.cell(0, i)
        cell.text = table_regular[i]
    for i in range(5):
        cell = b_c_table.cell(i + 1, 0)
        cell.text = str(i + 1)

def para_add(doc, item_para, item):
    if item == "货币资金":
        item_para.add_run("其中现金【】万元，银行存款【】万元、其他货币资金【】万元。")
    elif item == "应收票据":
        item_para.add_run("其中银行承兑汇票【】万元，商业承兑汇票【】万元。")
    elif item == "应收账款":
        item_para.add_run("账面余额【】万元、计提坏账准备【】万元，账龄1年以内占比【】%，3年以上占比【】%。其中，应收账款前五位：")
        item_table(doc, "序号", "名称", "余额（万元）", "占比")
    elif item == "预付账款":
        item_para.add_run("账龄1年以内占比【】%，3年以上占比【】%。其中，预收账款前五位：")
        item_table(doc, "序号", "名称", "余额（万元）", "占比")
    elif item == "其他应收款":
        item_para.add_run("账面余额【】万元、计提坏账准备【】万元。其中，其他应付款前五位：")
        item_table(doc, "序号", "名称", "余额(万元)", "款项性质")
    elif item == "存货":
        item_para.add_run("其中，原材料【】万元、库存商品【】万元、周转材料【】万元、工程施工【】万元、开发成本【】万元。")
    elif item == "其他流动资产":
        item_para.add_run("其中【添加明细】。")
    elif item == "长期股权投资":
        item_para.add_run("系对【N】家企业的投资，本期主要新增【哪家公司】；对外投资前五位如下：")
        item_table(doc, "序号", "名称", "投资额", "投资性质")
    elif item == "固定资产":
        item_para.add_run("固定资产原值【】万元，累计折旧【】万元，其中房屋及建筑物【】万元、机器设备【】万元、办公设备【】万元、【其他】【】万元。")
    elif item == "在建工程":
        item_para.add_run("主要为【项目1】【】万元、【项目2】【】万元、【项目3】【】万元……")
    elif item == "无形资产":
        item_para.add_run("主要为土地使用权【】万元、采矿权【】万元、专利权【】万元、软件【】万元，其他【】万元。")
    elif item == "短期借款":
        item_para.add_run("主要为【XX银行】【】万元、【XX银行】【】万元、【XX银行】【】万元、【xx银行】【】万元、【xx银行】【】万元。【其他需要说明的内容】")
    elif item == "应付票据":
        item_para.add_run("主要为银行承兑汇票【】万元，商业承兑汇票【】万元。")
    elif item == "应付账款":
        item_para.add_run("其中应付材料款【】万元，应付工程款【】万元。其中前五名如下：")
        item_table(doc, "序号", "名称", "余额", "性质")
    elif item == "预收账款":
        item_para.add_run("其中前5名如下：")
        item_table(doc, "序号", "名称", "余额", "账龄")
    elif item == "其他应付款":
        item_para.add_run("应付利息【】万元，往来款【】万元，押金和保证金【】万元，其中前5名如下：")
        item_table(doc, "序号", "名称", "余额", "账龄")
    elif item == "其他流动负债":
        item_para.add_run("主要为【】")
    elif item == "长期借款":
        item_para.add_run("主要为【XX银行】【】万元、【XX银行】【】万元、【XX银行】【】万元、【xx银行】【】万元、【xx银行】【】万元。【其他需要说明的内容】")
    elif item == "应付债券":
        item_para.add_run("主要为【】")
        item_table(doc, "序号", "债券名称", "余额", "到期日")
    elif item == "长期应付款":
        item_para.add_run("其中专项应付款【】万元、【】【】万元、其他【】万元。")

def items_detail(conf_obj, doc, date):

    #变量声明
    list_dict = []
    item_text = ''
    type_in_all = ''
    total_ = float()
    # 形成科目列表
    for dict in conf_obj.table_for_print:
        try:
            if data(conf_obj,dict['items'],'type') in ['流动资产', '非流动资产', '流动负债', '非流动负债']:
                list_dict.append(dict['items'])
            else:
                pass
        except Exception as Er:
            pass
    # 科目分析
    for item in list_dict:  # 第一行是表头，里面有字符串，必须剔除
        # 设置科目分析的文字模版
        if data(conf_obj, item, date) != 0:
            if conf_obj.data_type == 'no_1year':
                item_text = '【{}】：当期余额{:,.2f}万元，在{}中占比{:.2%}。'
            else:
                item_text = '【{}】：当期余额{:,.2f}万元，在{}中占比{:.2%}，较上年增加{:.2f}万元，{}。'

            if data(conf_obj, item, 'type') in ['流动资产', '非流动资产']:
                type_in_all = '总资产'
                total_ = data(conf_obj, '资产总计', date)
            elif data(conf_obj, item, 'type') in ['流动负债', '非流动负债']:
                type_in_all = '总负债'
                total_ = data(conf_obj, '负债合计', date)
            else:
                pass
        else:
            pass
        # “较上年增加后”的文字部分
        if isinstance(data(conf_obj, item, 'ratio'), str):
            text_ratio = '为' + data(conf_obj, item, 'ratio')
        else:
            text_ratio = '当年增幅为{:.2%}'.format(data(conf_obj, item, 'ratio'))
        # 模版的格式化输出
        if conf_obj.data_type == 'no_1year':
            para_ = doc.add_paragraph(
                item_text.format(
                    data(conf_obj, item,'items'),
                    data(conf_obj, item, date),
                    type_in_all,
                    data(conf_obj, item, date) / total_
                )
            )
            para_.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_.paragraph_format.first_line_indent = Pt(24)
            para_add(doc, para_, data(conf_obj, item, date))  # 根据不同的科目在文字模版后新增文字
        else:
            para_ = doc.add_paragraph(
                item_text.format(
                    data(conf_obj, item, 'items'),
                    data(conf_obj, item, date),
                    type_in_all,
                    data(conf_obj, item, date) / total_,
                    data(conf_obj, item, 'delta'),
                    text_ratio
                )
            )
            para_.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            para_.paragraph_format.first_line_indent = Pt(24)
            para_add(doc, para_, data(conf_obj, item, 'items')) #根据不同的科目在文字模版后新增文字
# ----------------------------------------------



